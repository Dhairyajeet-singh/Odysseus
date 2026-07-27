"""Generate test resumes that reproduce each failure mode.

Each fixture targets one thing the old pipeline got wrong:

  single_column.pdf   baseline — must not regress
  two_column.pdf      sidebar + main column: the silent line-fusion bug
  table_skills.pdf    skills matrix in a ruled table
  scanned.pdf         image-only, no text layer at all
  junk_layer.pdf      image + 130-char watermark text layer -> defeats `len < 50`
  designer.docx       DOCX with text box, table, header and hyperlink
  two_column_copy.pdf near-duplicate of two_column.pdf (re-export simulation)
"""

from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from docx.oxml import parse_xml
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

OUT = Path(__file__).parent / "fixtures"
W, H = A4  # 595 x 842 pt

SIDEBAR = [
    ("SKILLS", 11, True),
    ("Python", 9, False), ("SQL", 9, False), ("Docker", 9, False),
    ("Kubernetes", 9, False), ("PostgreSQL", 9, False), ("Airflow", 9, False),
    ("PyTorch", 9, False), ("Terraform", 9, False),
    ("", 9, False),
    ("EDUCATION", 11, True),
    ("B.Tech, Computer Science", 9, False),
    ("NIT Trichy, 2018", 9, False),
    ("CGPA 8.7 / 10", 9, False),
    ("", 9, False),
    ("LANGUAGES", 11, True),
    ("English, Hindi, Tamil", 9, False),
]

MAIN = [
    ("EXPERIENCE", 12, True),
    ("Senior Data Engineer, Acme Corp", 10, True),
    ("2021 - 2024, Bengaluru", 9, False),
    ("Built a streaming ingestion platform handling 2M events", 9, False),
    ("per minute using Kafka and Flink, cutting pipeline latency", 9, False),
    ("from 40 minutes to under 90 seconds.", 9, False),
    ("Led migration of 60 batch jobs from Airflow 1 to Airflow 2", 9, False),
    ("with zero downtime across a six week window.", 9, False),
    ("", 9, False),
    ("Data Engineer, Beta Analytics", 10, True),
    ("2018 - 2021, Chennai", 9, False),
    ("Designed the dimensional warehouse backing all executive", 9, False),
    ("reporting, reducing dashboard query time by 70 percent.", 9, False),
    ("Owned data quality tooling used by 40 analysts daily.", 9, False),
    ("", 9, False),
    ("PROJECTS", 12, True),
    ("Resume ranking service, personal project", 10, True),
    ("Semantic matching over 12k resumes using sentence", 9, False),
    ("transformers with a calibrated scoring layer.", 9, False),
]


def _header(c: canvas.Canvas, name: str = "PRIYA RAMACHANDRAN") -> float:
    c.setFont("Helvetica-Bold", 20)
    c.drawCentredString(W / 2, H - 60, name)
    c.setFont("Helvetica", 9)
    c.drawCentredString(
        W / 2, H - 78,
        "priya.ramachandran@example.com  |  +91 98765 43210  |  "
        "linkedin.com/in/priyar  |  Bengaluru, India",
    )
    c.setLineWidth(0.7)
    c.line(40, H - 90, W - 40, H - 90)
    return H - 120


def two_column(path: Path, name: str = "PRIYA RAMACHANDRAN") -> None:
    """Sidebar at x=40..185, gutter 185..235, main column at x=235..555."""
    c = canvas.Canvas(str(path), pagesize=A4)
    top = _header(c, name)

    y = top
    for text, size, bold in SIDEBAR:
        if text:
            c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
            c.drawString(40, y, text)
        y -= size + 5

    y = top
    for text, size, bold in MAIN:
        if text:
            c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
            c.drawString(235, y, text)
        y -= size + 5

    c.showPage()
    c.save()


def single_column(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    y = _header(c, "ARJUN MEHTA")
    for text, size, bold in MAIN + [("", 9, False)] + SIDEBAR:
        if text:
            c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
            c.drawString(50, y, text)
        y -= size + 5
    c.showPage()
    c.save()


def table_skills(path: Path) -> None:
    """Skills as a ruled grid — the layout that scrambles under naive parsing."""
    c = canvas.Canvas(str(path), pagesize=A4)
    y = _header(c, "SANA KHURANA")

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "TECHNICAL SKILLS")
    y -= 22

    rows = [
        ["Category", "Primary", "Secondary"],
        ["Languages", "Python, Go", "Java, Rust"],
        ["Cloud", "AWS, GCP", "Azure"],
        ["Data", "Spark, dbt", "Snowflake"],
        ["ML", "PyTorch", "scikit-learn"],
    ]
    x0, col_w, row_h = 50, 150, 20
    for r, row in enumerate(rows):
        for col, cell in enumerate(row):
            x = x0 + col * col_w
            c.rect(x, y - row_h, col_w, row_h, stroke=1, fill=0)
            c.setFont("Helvetica-Bold" if r == 0 else "Helvetica", 9)
            c.drawString(x + 5, y - 14, cell)
        y -= row_h

    y -= 30
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "EXPERIENCE")
    y -= 18
    c.setFont("Helvetica", 9)
    for line in ["Machine Learning Engineer, Delta Labs (2020 - 2024)",
                 "Shipped a retrieval pipeline serving 400 QPS at p99 under 120ms.",
                 "Owned model evaluation harness and offline/online metric parity."]:
        c.drawString(50, y, line)
        y -= 14
    c.showPage()
    c.save()


def _rasterise(src: Path, dpi: int = 200) -> bytes:
    doc = fitz.open(str(src))
    pix = doc[0].get_pixmap(dpi=dpi)
    png = pix.tobytes("png")
    doc.close()
    return png


def scanned(path: Path, src: Path) -> None:
    """Image-only PDF: exactly what a flatbed scan or a phone photo produces."""
    png = _rasterise(src)
    doc = fitz.open()
    page = doc.new_page(width=W, height=H)
    page.insert_image(fitz.Rect(0, 0, W, H), stream=png)
    doc.save(str(path))
    doc.close()


def junk_layer(path: Path, src: Path) -> None:
    """The case `len(text) < 50` cannot catch.

    A scanned page carrying a short, real text layer — an export watermark and
    a page footer. 130+ characters of perfectly valid text sitting on top of an
    image that holds the actual resume. The old heuristic sees >50 chars, skips
    OCR, and ranks the candidate on a watermark.
    """
    png = _rasterise(src)
    doc = fitz.open()
    page = doc.new_page(width=W, height=H)
    page.insert_image(fitz.Rect(0, 0, W, H), stream=png)
    page.insert_text((40, 812),
                     "Generated by ScanPro 4.2 - www.scanpro.example.com",
                     fontsize=7, color=(0.6, 0.6, 0.6))
    page.insert_text((40, 824),
                     "Document reference SCN-2024-88213 - Page 1 of 1 - Confidential",
                     fontsize=7, color=(0.6, 0.6, 0.6))
    doc.save(str(path))
    doc.close()


_TEXTBOX_XML = """
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
     xmlns:v="urn:schemas-microsoft-com:vml">
  <w:r>
    <mc:AlternateContent>
      <mc:Choice Requires="wps">
        <w:drawing><wps:txbx><w:txbxContent>
          <w:p><w:r><w:t>CORE SKILLS</w:t></w:r></w:p>
          <w:p><w:r><w:t>Python, Django, FastAPI, Redis, Celery</w:t></w:r></w:p>
          <w:p><w:r><w:t>AWS Lambda, DynamoDB, CloudFormation</w:t></w:r></w:p>
        </w:txbxContent></wps:txbx></w:drawing>
      </mc:Choice>
      <mc:Fallback>
        <w:pict><v:shape><v:textbox><w:txbxContent>
          <w:p><w:r><w:t>CORE SKILLS</w:t></w:r></w:p>
          <w:p><w:r><w:t>Python, Django, FastAPI, Redis, Celery</w:t></w:r></w:p>
          <w:p><w:r><w:t>AWS Lambda, DynamoDB, CloudFormation</w:t></w:r></w:p>
        </w:txbxContent></v:textbox></v:shape></w:pict>
      </mc:Fallback>
    </mc:AlternateContent>
  </w:r>
</w:p>
"""


def designer_docx(path: Path) -> None:
    """DOCX with the four things naive `.paragraphs` iteration misses."""
    doc = Document()

    doc.sections[0].header.paragraphs[0].text = (
        "rohan.desai@example.com | +91 91234 56789"
    )

    doc.add_heading("Rohan Desai", level=0)
    doc.add_paragraph("Backend engineer, six years building payment systems.")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Staff Engineer, Fintech Co (2020 - 2024)")
    doc.add_paragraph("Rebuilt the settlement engine; cut reconciliation "
                      "breaks by 92 percent.")

    # Text box (Word emits it twice — mc:Choice and mc:Fallback).
    doc.element.body.append(parse_xml(_TEXTBOX_XML))

    doc.add_heading("Education", level=1)
    t = doc.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    data = [["Degree", "Institution", "Year"],
            ["M.Tech CSE", "IIT Bombay", "2018"],
            ["B.E. IT", "Pune University", "2016"]]
    for r, row in enumerate(data):
        for col, val in enumerate(row):
            t.cell(r, col).text = val

    # Hyperlink whose visible text hides the URL.
    p = doc.add_paragraph()
    part = doc.part
    r_id = part.relate_to(
        "https://github.com/rohandesai",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    p._p.append(parse_xml(
        f'<w:hyperlink xmlns:w="{_TEXTBOX_XML.split(chr(34))[1]}" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/'
        f'relationships" r:id="{r_id}"><w:r><w:t>GitHub</w:t></w:r></w:hyperlink>'
    ))

    doc.save(str(path))


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    single_column(OUT / "single_column.pdf")
    two_column(OUT / "two_column.pdf")
    two_column(OUT / "two_column_copy.pdf", name="PRIYA  RAMACHANDRAN")
    table_skills(OUT / "table_skills.pdf")
    scanned(OUT / "scanned.pdf", OUT / "two_column.pdf")
    junk_layer(OUT / "junk_layer.pdf", OUT / "two_column.pdf")
    designer_docx(OUT / "designer.docx")
    for f in sorted(OUT.iterdir()):
        print(f"  {f.name:24s} {f.stat().st_size:>8,} bytes")


if __name__ == "__main__":
    main()
