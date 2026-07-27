"""DOCX extraction.

The assignment asks for PDF **and** DOCX. The naive approach --
`"\\n".join(p.text for p in Document(path).paragraphs)` -- silently loses four
things that matter on real resumes:

* **Tables.** `document.paragraphs` skips them entirely. Plenty of resumes are
  built entirely inside one invisible table, which under that approach extracts
  to an empty string.
* **Text boxes.** Designer templates put the whole sidebar -- skills, contact
  details -- inside a text box. `paragraphs` never sees it.
* **Headers/footers.** Contact details live there more often than you'd expect.
* **Hyperlink targets.** The GitHub URL exists only as a relationship target;
  the visible text is often just "GitHub".

So we walk the OOXML body in document order instead, which also keeps tables
positioned correctly relative to surrounding paragraphs rather than appended at
the end.

Positions are synthetic: DOCX is a flow format with no fixed geometry until it
is laid out. Blocks get monotonically increasing pseudo-y values so they share
the Block type with the PDF path, but no column detection is attempted -- there
are no columns to detect until rendering.
"""

from __future__ import annotations

import zipfile
from typing import List, Tuple

from docx import Document
from docx.table import Table as DocxTable

from .types import Block, Table

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"

_TXBX = f"{_W}txbxContent"
_FALLBACK = f"{_MC}Fallback"


def _iter_text(elem, skip: Tuple[str, ...] = (_TXBX, _FALLBACK)) -> List[str]:
    """Depth-first `w:t` collection, skipping named subtrees.

    Skipping `mc:Fallback` matters: modern Word writes every text box twice --
    once under `mc:Choice` as a DrawingML shape and once under `mc:Fallback` as
    a legacy VML shape. Collecting both duplicates the entire sidebar, which
    would double-count every skill in it.
    """
    out: List[str] = []
    for child in elem:
        if child.tag in skip:
            continue
        if child.tag == f"{_W}t" and child.text:
            out.append(child.text)
        elif child.tag == f"{_W}tab":
            out.append("\t")
        elif child.tag in (f"{_W}br", f"{_W}cr"):
            out.append("\n")
        else:
            out.extend(_iter_text(child, skip))
    return out


def _para_text(p) -> str:
    return "".join(_iter_text(p)).strip()


def _textbox_paras(p) -> List[str]:
    """Paragraph texts from text boxes anchored in this paragraph."""
    out: List[str] = []
    for tb in p.iter(_TXBX):
        # Skip the VML duplicate under mc:Fallback.
        parent = tb.getparent()
        legacy = False
        while parent is not None:
            if parent.tag == _FALLBACK:
                legacy = True
                break
            parent = parent.getparent()
        if legacy:
            continue
        for tp in tb.iter(f"{_W}p"):
            t = "".join(_iter_text(tp, skip=())).strip()
            if t:
                out.append(t)
    return out


def _para_style(p) -> Tuple[bool, float | None]:
    """(is_bold, size_pt) from the paragraph's first run — used for headings."""
    bold = False
    size = None
    for rpr in p.iter(f"{_W}rPr"):
        if rpr.find(f"{_W}b") is not None:
            bold = True
        sz = rpr.find(f"{_W}sz")
        if sz is not None and size is None:
            try:
                size = float(sz.get(f"{_W}val")) / 2.0  # half-points
            except (TypeError, ValueError):
                pass
        break
    for pstyle in p.iter(f"{_W}pStyle"):
        val = (pstyle.get(f"{_W}val") or "").lower()
        if val.startswith("heading") or val == "title":
            bold = True
        break
    return bold, size


def _hyperlinks(doc) -> List[str]:
    out: List[str] = []
    try:
        rels = doc.part.rels
        for rel in rels.values():
            if "hyperlink" in rel.reltype and rel.is_external:
                if rel.target_ref not in out:
                    out.append(rel.target_ref)
    except Exception:
        pass
    return out


def _headers_footers(doc) -> List[str]:
    out: List[str] = []
    for section in doc.sections:
        for part in (section.header, section.footer):
            try:
                for p in part.paragraphs:
                    t = p.text.strip()
                    if t and t not in out:
                        out.append(t)
            except Exception:
                continue
    return out


def is_docx(path: str) -> bool:
    try:
        with zipfile.ZipFile(path) as z:
            return "word/document.xml" in z.namelist()
    except Exception:
        return False


def extract(path: str) -> Tuple[List[Block], List[Table], List[str], List[str]]:
    """Returns (blocks, tables, links, warnings)."""
    warnings: List[str] = []
    doc = Document(path)
    body = doc.element.body

    blocks: List[Block] = []
    tables: List[Table] = []
    y = 0.0

    def add(text: str, source: str, *, bold: bool = False,
            size: float | None = None, in_table: bool = False) -> None:
        nonlocal y
        if not text.strip():
            return
        y += 12.0
        blocks.append(Block(
            text=text.strip(), page=1, bbox=(0.0, y, 468.0, y + 10.0),
            source=source, font_size=size, bold=bold, in_table=in_table,
        ))

    for child in body.iterchildren():
        if child.tag == f"{_W}p":
            bold, size = _para_style(child)
            add(_para_text(child), "docx:body", bold=bold, size=size)
            for tb_text in _textbox_paras(child):
                add(tb_text, "docx:textbox")

        elif child.tag == f"{_W}tbl":
            try:
                t = DocxTable(child, doc)
                rows = [[(c.text or "").strip() for c in row.cells] for row in t.rows]
                rows = [r for r in rows if any(r)]
            except Exception as exc:  # malformed table markup
                warnings.append(f"table skipped: {exc}")
                continue
            if not rows:
                continue
            if len(rows) >= 2 and len(rows[0]) >= 2:
                tables.append(Table(page=1, bbox=None, rows=rows, source="docx"))
            # Emit cell text in reading order too, so single-column layout
            # tables (very common as an invisible resume grid) still produce
            # coherent prose rather than only a structured table object.
            for r in rows:
                add(" | ".join(c for c in r if c), "docx:table", in_table=True)

    hf = _headers_footers(doc)
    if hf:
        add("\n".join(hf), "docx:header_footer")

    links = _hyperlinks(doc)

    if not blocks:
        warnings.append(
            "no text found in DOCX — document may contain only images "
            "(convert to PDF and run the OCR path)"
        )
    return blocks, tables, links, warnings
